"""Generate a branded PDF welcome letter for one or more drivers.

Each driver gets a full-page welcome letter produced from the DOCX
marketing template (``driver_welcome_letter_template.docx``).  The DOCX
is filled with driver-specific data via python-docx, then converted to
PDF through LibreOffice headless — preserving fonts, images, and layout
exactly as designed.

Company details (name, website, email) are pulled from admin settings so
the letter stays current without code changes.
"""

from __future__ import annotations

import logging
import os
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore[import-untyped]

logger = logging.getLogger(__name__)

_TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "static" / "branding" / "driver_welcome_letter_template.docx"


def _driver_name(driver: dict[str, Any]) -> str:
    first = driver.get("first_name") or ""
    last = driver.get("last_name") or ""
    return f"{first} {last}".strip() or driver.get("name", "Driver")


def _driver_address(driver: dict[str, Any]) -> str:
    city = driver.get("service_area_name") or driver.get("city") or ""
    meta = driver.get("legacy_import_metadata") or {}
    if isinstance(meta, dict) and meta.get("address"):
        return str(meta["address"])
    if city:
        return city
    return ""


def _replace_in_runs(paragraph, old: str, new: str) -> None:
    """Replace ``old`` with ``new`` across the runs of a paragraph.

    Handles both the simple case (placeholder fully inside one run) and
    the split case (placeholder fragmented across adjacent runs by Word's
    internal markup decisions).
    """
    full_text = "".join(r.text for r in paragraph.runs)
    if old not in full_text:
        return

    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return

    # Placeholder is split across runs — rebuild
    combined = ""
    start_idx = None
    for i, run in enumerate(paragraph.runs):
        combined += run.text
        if start_idx is None and old[0] in run.text:
            start_idx = i
        if old in combined and start_idx is not None:
            prefix = combined[: combined.index(old)]
            suffix = combined[combined.index(old) + len(old) :]
            paragraph.runs[start_idx].text = prefix + new + suffix
            for j in range(start_idx + 1, i + 1):
                paragraph.runs[j].text = ""
            return


def _replace_in_table_cells(table, old: str, new: str) -> None:
    """Replace ``old`` with ``new`` in every cell of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_runs(paragraph, old, new)


def _fill_template(
    driver: dict[str, Any],
    company: dict[str, Any] | None = None,
) -> bytes:
    """Fill the DOCX template for one driver and return raw DOCX bytes."""
    company = company or {}
    company_name = company.get("company_name") or "Spinr"
    website = company.get("company_website") or "www.spinr.ca"
    contact_email = company.get("company_email") or "drivers@spinr.ca"

    if website.startswith("https://"):
        website_display = website[8:]
    elif website.startswith("http://"):
        website_display = website[7:]
    else:
        website_display = website

    name = _driver_name(driver)
    address = _driver_address(driver)

    doc = Document(str(_TEMPLATE_PATH))

    replacements = {
        "«Name»": name,
        "«Address»": address,
    }
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            _replace_in_runs(paragraph, old, new)
        if company_name != "Spinr":
            _replace_in_runs(paragraph, "Spinr", company_name)
        if website_display != "www.spinr.ca":
            _replace_in_runs(paragraph, "www.spinr.ca", website_display)

    for table in doc.tables:
        if contact_email != "drivers@spinr.ca":
            _replace_in_table_cells(table, "drivers@spinr.ca", contact_email)
        if company_name != "Spinr":
            _replace_in_table_cells(table, "Spinr", company_name)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_to_pdf(docx_bytes: bytes, work_dir: str) -> bytes:
    """Convert a DOCX byte-string to PDF via LibreOffice headless."""
    docx_path = os.path.join(work_dir, "letter.docx")
    with open(docx_path, "wb") as f:
        f.write(docx_bytes)

    result = subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--norestore",
            "--convert-to",
            "pdf",
            "--outdir",
            work_dir,
            docx_path,
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        logger.error("LibreOffice conversion failed: %s", result.stderr)
        raise RuntimeError(f"LibreOffice PDF conversion failed: {result.stderr}")

    pdf_path = os.path.join(work_dir, "letter.pdf")
    if not os.path.exists(pdf_path):
        raise RuntimeError("LibreOffice did not produce a PDF file")

    with open(pdf_path, "rb") as f:
        return f.read()


def generate_decal_pdf(
    drivers: list[dict[str, Any]],
    *,
    company: dict[str, Any] | None = None,
) -> bytes:
    """Generate a multi-page PDF with one welcome letter per driver.

    Returns the raw PDF bytes suitable for a Response(content=...,
    media_type="application/pdf").

    Each driver's letter is produced by filling the DOCX marketing
    template with their name/address, then converting to PDF via
    LibreOffice. Multiple drivers are merged into a single PDF using
    pypdf.
    """
    from pypdf import PdfReader, PdfWriter  # type: ignore[import-untyped]

    writer = PdfWriter()

    for driver in drivers:
        work_dir = tempfile.mkdtemp(prefix="spinr_decal_")
        try:
            docx_bytes = _fill_template(driver, company=company)
            pdf_bytes = _docx_to_pdf(docx_bytes, work_dir)
            reader = PdfReader(BytesIO(pdf_bytes))
            for page in reader.pages:
                writer.add_page(page)
        finally:
            shutil.rmtree(work_dir, ignore_errors=True)

    output = BytesIO()
    writer.write(output)
    return output.getvalue()
